#!/usr/bin/env python3
"""Render the reviewed Markdown framework into a polished DOCX.

The Markdown remains the semantic source of truth. This builder creates the
editable/printable delivery view using the compact_reference_guide preset.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "自动驾驶单次运行实时性—物理安全双向诊断框架.md"
OUTPUT = ROOT / "自动驾驶单次运行实时性—物理安全双向诊断框架.docx"

# compact_reference_guide tokens
PAGE_W_IN = 8.5
PAGE_H_IN = 11.0
MARGIN_IN = 1.0
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS = {"top": 80, "bottom": 80, "start": 120, "end": 120}
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "667085"
TABLE_FILL = "E8EEF5"
LIGHT_FILL = "F4F6F9"
RULE = "C9D5E2"
WHITE = "FFFFFF"
RISK = "9B1C1C"
BODY_FONT = "Arial Unicode MS"
MONO_FONT = "Arial Unicode MS"


def set_run_font(run, *, name=BODY_FONT, east_asia=BODY_FONT, size=None,
                 color=None, bold=None, italic=None):
    run.font.name = name
    rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_margins(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin in ("top", "start", "bottom", "end"):
        if margin in kwargs:
            node = tc_mar.find(qn(f"w:{margin}"))
            if node is None:
                node = OxmlElement(f"w:{margin}")
                tc_mar.append(node)
            node.set(qn("w:w"), str(kwargs[margin]))
            node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color=RULE, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:color"), color)
        el.set(qn("w:space"), "0")


def set_table_geometry(table, widths_dxa):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell, **CELL_MARGINS)


def paragraph_shading(paragraph, fill, border=None):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    if border:
        p_bdr = p_pr.find(qn("w:pBdr"))
        if p_bdr is None:
            p_bdr = OxmlElement("w:pBdr")
            p_pr.append(p_bdr)
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:color"), border)
        left.set(qn("w:space"), "8")
        p_bdr.append(left)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    tail = paragraph.add_run(" 页")
    set_run_font(tail, size=9, color=MUTED)


def add_hyperlink(paragraph, text, url, color=BLUE):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    r_pr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), BODY_FONT)
    rfonts.set(qn("w:hAnsi"), BODY_FONT)
    rfonts.set(qn("w:eastAsia"), BODY_FONT)
    r_pr.append(rfonts)
    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_RE = re.compile(
    r"(\\\(.+?\\\)|\*\*.+?\*\*|`.+?`|\[[^\]]+\]\(https?://[^)]+\))"
)


def add_inline(paragraph, text, *, size=None, color=None, base_bold=False):
    pos = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_run_font(run, size=size, color=color, bold=base_bold)
        token = match.group(0)
        if token.startswith(r"\("):
            run = paragraph.add_run(linearize_latex(token[2:-2]))
            set_run_font(run, size=size, color=color, bold=base_bold)
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, color=color, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name=MONO_FONT, east_asia=BODY_FONT,
                         size=(size or 11) - 0.5, color=INK, bold=base_bold)
            run._element.get_or_add_rPr().append(OxmlElement("w:noProof"))
        else:
            m = re.match(r"\[([^\]]+)\]\((https?://[^)]+)\)", token)
            add_hyperlink(paragraph, m.group(1), m.group(2))
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=size, color=color, bold=base_bold)


LATEX_MAP = {
    r"\tau": "τ", r"\rho": "ρ", r"\Phi": "Φ", r"\Theta": "Θ",
    r"\Delta": "Δ", r"\mathbf": "", r"\boldsymbol": "",
    r"\rightarrow": "→", r"\leftarrow": "←", r"\mid": "|",
    r"\parallel": "∥", r"\ge": "≥", r"\le": "≤", r"\in": "∈",
    r"\sup": "sup", r"\inf": "inf", r"\min": "min", r"\max": "max",
    r"\int": "∫", r"\frac": "frac", r"\boxed": "",
    r"\begin{bmatrix}": "[", r"\end{bmatrix}": "]",
    r"\begin{aligned}": "", r"\end{aligned}": "",
    r"\text": "", r"\underbrace": "", r"\quad": "  ",
    r"\left": "", r"\right": "", r"\;": " ", r"\,": " ",
    r"\cdot": "·", r"\ldots": "…", r"\Vert": "‖",
    r"\sum": "Σ", r"\sqrt": "sqrt", r"\infty": "∞",
    r"\forall": "∀", r"\mu": "μ", r"\theta": "θ",
    r"\partial": "∂", r"\nabla": "∇", r"\approx": "≈",
    r"\neq": "≠", r"\times": "×", r"\pm": "±",
    r"\mathrm": "", r"\operatorname": "", r"\mathcal": "",
    r"\dot": "d/dt ", r"\!": "",
}


def linearize_latex(text):
    value = " ".join(line.strip() for line in text.splitlines())
    # Longest-first avoids prefix collisions such as ``\\in`` corrupting
    # ``\\inf`` and ``\\left`` corrupting ``\\leftarrow``.
    for src, dst in sorted(LATEX_MAP.items(), key=lambda item: len(item[0]), reverse=True):
        value = value.replace(src, dst)
    value = value.replace(r"\\", "; ")
    # Make simple fractions legible: frac{a}{b} -> (a)/(b)
    for _ in range(4):
        value = re.sub(r"frac\{([^{}]+)\}\{([^{}]+)\}", r"(\1)/(\2)", value)
    value = re.sub(r"\{([^{}]+)\}", r"\1", value)
    value = value.replace(r"\{", "{").replace(r"\}", "}")
    # Remove any remaining LaTeX control words so the printable view never
    # exposes renderer-specific commands such as ``\left`` or ``\sf``.
    value = re.sub(r"\\[A-Za-z]+", "", value)
    value = re.sub(r"\s+", " ", value).strip()
    return value


def add_equation(doc, latex):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.25)
    paragraph_shading(p, LIGHT_FILL)
    run = p.add_run(linearize_latex(latex))
    set_run_font(run, name=BODY_FONT, east_asia=BODY_FONT,
                 size=10.5, color=INK)
    return p


def add_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(7)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), RULE)
    bottom.set(qn("w:space"), "1")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def parse_table(lines, start):
    rows = []
    idx = start
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        raw = lines[idx].strip().strip("|")
        rows.append([c.strip() for c in raw.split("|")])
        idx += 1
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in rows[1]):
        rows.pop(1)
    return rows, idx


def column_widths(rows):
    n = max(len(r) for r in rows)
    scores = []
    for col in range(n):
        longest = max((len(r[col]) if col < len(r) else 0) for r in rows)
        scores.append(max(8, min(longest, 45)))
    total_score = sum(scores)
    widths = [int(CONTENT_DXA * s / total_score) for s in scores]
    widths[-1] += CONTENT_DXA - sum(widths)
    return widths


def add_markdown_table(doc, rows):
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    widths = column_widths(rows)
    set_table_geometry(table, widths)
    set_table_borders(table)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for r_idx, row in enumerate(rows):
        for c_idx in range(ncols):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if r_idx == 0:
                shade_cell(cell, TABLE_FILL)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.12
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            text = row[c_idx] if c_idx < len(row) else ""
            add_inline(p, text, size=9.2, color=(INK if r_idx == 0 else None), base_bold=(r_idx == 0))
    set_repeat_table_header(table.rows[0])
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(0)
    after.paragraph_format.space_after = Pt(2)
    return table


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(PAGE_W_IN)
    section.page_height = Inches(PAGE_H_IN)
    section.top_margin = Inches(MARGIN_IN)
    section.bottom_margin = Inches(MARGIN_IN)
    section.left_margin = Inches(MARGIN_IN)
    section.right_margin = Inches(MARGIN_IN)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    p = header.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("TCPS-PA v3.1  ·  单次运行实时性—物理安全双向诊断框架")
    set_run_font(run, size=8.5, color=MUTED, bold=True)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:color"), RULE)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    footer = section.footer
    add_page_field(footer.paragraphs[0])


def add_cover(doc):
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(16)
    r = kicker.add_run("METHOD REFERENCE  ·  TCPS-PA v3.1")
    set_run_font(r, size=10, color=BLUE, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    title.paragraph_format.keep_with_next = True
    r = title.add_run("自动驾驶单次运行实时性—物理安全\n双向诊断框架")
    set_run_font(r, size=28, color=INK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(28)
    r = subtitle.add_run("Event-Centric · Dynamic Timing Contract · Bidirectional Diagnosis · Physical Budget Loss")
    set_run_font(r, size=12.5, color=DARK_BLUE)

    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = rule.add_run("━" * 32)
    set_run_font(r, size=8, color=RULE)
    rule.paragraph_format.space_after = Pt(20)

    for label, value in (
        ("适用系统", "CARLA 0.9.15 + Apollo 10.0.0 + Bridge"),
        ("分析范围", "一次 run 内的发现、定位、判定、解释与物理损失量化"),
        ("证据约束", "Observed / Requirement / Model / Retrospective 严格分离"),
        ("版本日期", "2026-08-13"),
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(f"{label}  ")
        set_run_font(r, size=10.5, color=MUTED, bold=True)
        r = p.add_run(value)
        set_run_font(r, size=10.5, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(44)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("本方法只针对单次运行；不包含多 run 组间比较方法")
    set_run_font(r, size=10, color=RISK, bold=True)
    doc.add_page_break()


def render_markdown(doc, text):
    lines = text.splitlines()
    # Skip source title/metadata block through first horizontal rule; cover replaces it.
    idx = 0
    while idx < len(lines) and lines[idx].strip() != "---":
        idx += 1
    idx += 1
    in_code = False
    code_lines = []
    in_math = False
    math_lines = []

    while idx < len(lines):
        raw = lines[idx]
        stripped = raw.strip()
        if in_code:
            if stripped.startswith("```"):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.18)
                p.paragraph_format.right_indent = Inches(0.18)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(7)
                p.paragraph_format.line_spacing = 1.08
                paragraph_shading(p, LIGHT_FILL, border=RULE)
                r = p.add_run("\n".join(code_lines))
                set_run_font(r, name=MONO_FONT, east_asia=BODY_FONT, size=8.8, color=INK)
                code_lines = []
                in_code = False
            else:
                code_lines.append(raw)
            idx += 1
            continue
        if in_math:
            if stripped == r"\]":
                add_equation(doc, "\n".join(math_lines))
                math_lines = []
                in_math = False
            else:
                math_lines.append(raw)
            idx += 1
            continue
        if stripped.startswith("```"):
            in_code = True
            idx += 1
            continue
        if stripped == r"\[":
            in_math = True
            idx += 1
            continue
        if not stripped:
            idx += 1
            continue
        if stripped == "---":
            add_rule(doc)
            idx += 1
            continue
        if stripped.startswith("|") and idx + 1 < len(lines) and lines[idx + 1].strip().startswith("|"):
            rows, idx = parse_table(lines, idx)
            add_markdown_table(doc, rows)
            continue
        m = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if m:
            # The source H1 is replaced by the cover. Promote source H2/H3 so
            # the DOCX has a real Heading 1/2 hierarchy for navigation/TOC.
            level = max(1, len(m.group(1)) - 1)
            p = doc.add_paragraph(style=f"Heading {level}")
            add_inline(p, m.group(2))
            idx += 1
            continue
        if stripped.startswith(">"):
            quote_lines = []
            while idx < len(lines) and lines[idx].strip().startswith(">"):
                quote_lines.append(lines[idx].strip()[1:].strip())
                idx += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.right_indent = Inches(0.12)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(7)
            paragraph_shading(p, LIGHT_FILL, border=BLUE)
            add_inline(p, " ".join(quote_lines), color=INK)
            continue
        m = re.match(r"^[-*]\s+(.+)$", stripped)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, m.group(1))
            idx += 1
            continue
        m = re.match(r"^\d+\.\s+(.+)$", stripped)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_inline(p, m.group(1))
            idx += 1
            continue

        para_lines = [stripped]
        idx += 1
        while idx < len(lines):
            nxt = lines[idx].strip()
            if (not nxt or nxt == "---" or nxt == r"\[" or nxt.startswith("```")
                    or nxt.startswith("#") or nxt.startswith("|") or nxt.startswith(">")
                    or re.match(r"^[-*]\s+", nxt) or re.match(r"^\d+\.\s+", nxt)):
                break
            para_lines.append(nxt)
            idx += 1
        p = doc.add_paragraph()
        add_inline(p, " ".join(para_lines))


def main():
    text = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    configure_styles(doc)
    configure_document(doc)
    add_cover(doc)
    render_markdown(doc, text)
    core = doc.core_properties
    core.title = "自动驾驶单次运行实时性—物理安全双向诊断框架"
    core.subject = "Single-run event-centric temporal correctness and physical safety diagnosis"
    core.author = "TCPS-PA Method Working Draft"
    core.keywords = "Apollo, CARLA, real-time systems, dynamic deadline, cause-effect chain"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
